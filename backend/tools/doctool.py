"""
Enhanced DOCX Export Tool with Multi-Level Support
Generates professional Word documents with level-appropriate formatting
"""
import base64
import io
import re
import json
import threading
from typing import Optional
from pydantic import BaseModel, Field
from langchain.tools import tool
from langchain_core.runnables import RunnableConfig

# Thread-safe storage for download data from subagent tool calls
# Uses a global latest-download approach instead of thread IDs because
# LangGraph runs tools via asyncio.run_in_executor() in a thread pool,
# so the storing thread ID differs from the retrieving thread ID.
_latest_download = None
_download_lock = threading.Lock()

def _store_pending_download(data: dict):
    """Store download data globally (latest wins)."""
    global _latest_download
    with _download_lock:
        _latest_download = data
        print(f"  💾 [doctool] Download stored in pending: {data.get('filename', 'unknown')}")

def get_pending_download() -> dict | None:
    """Retrieve and remove the latest pending download data."""
    global _latest_download
    with _download_lock:
        data = _latest_download
        _latest_download = None
        return data

import requests
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

class DocSection(BaseModel):
    """Represents a single section of the report."""
    heading: str = Field(description="The section title (e.g., 'Introduction')")
    content: str = Field(description="The full text content for this section (supports Markdown formatting)")

class ExportInput(BaseModel):
    """Input schema for the DOCX export tool."""
    sections: list[DocSection] = Field(description="A list of section objects")
    filename: str = Field(description="The output filename, e.g., 'report.docx'")
    report_level: str = Field(
        default="student",
        description="Report level: 'student', 'professor', or 'researcher'"
    )
    title: str = Field(default="Research Report", description="Document title for cover page")


def setup_styles(doc: Document, report_level: str):
    """
    Configure document styles based on report level.
    
    Student: Friendly, readable, clear hierarchy
    Professor: Professional, balanced, organized
    Researcher: Formal, academic, publication-ready
    """
    styles = doc.styles
    
    # Define color schemes by level
    level_colors = {
        "student": {
            "title": RGBColor(41, 128, 185),      # Friendly blue
            "heading1": RGBColor(52, 152, 219),   # Bright blue
            "heading2": RGBColor(46, 134, 193),   # Medium blue
            "accent": RGBColor(52, 152, 219)
        },
        "professor": {
            "title": RGBColor(44, 62, 80),        # Professional dark
            "heading1": RGBColor(52, 73, 94),     # Dark slate
            "heading2": RGBColor(69, 90, 100),    # Slate gray
            "accent": RGBColor(52, 73, 94)
        },
        "researcher": {
            "title": RGBColor(26, 35, 46),        # Academic dark
            "heading1": RGBColor(33, 33, 33),     # Almost black
            "heading2": RGBColor(66, 66, 66),     # Dark gray
            "accent": RGBColor(33, 33, 33)
        }
    }
    
    colors = level_colors.get(report_level, level_colors["student"])
    
    # Modify existing Heading 1 style
    heading1 = styles['Heading 1']
    heading1_font = heading1.font
    heading1_font.name = 'Arial'
    heading1_font.size = Pt(18 if report_level == "student" else 16)
    heading1_font.bold = True
    heading1_font.color.rgb = colors["heading1"]
    
    # Modify Heading 2 style
    heading2 = styles['Heading 2']
    heading2_font = heading2.font
    heading2_font.name = 'Arial'
    heading2_font.size = Pt(14)
    heading2_font.bold = True
    heading2_font.color.rgb = colors["heading2"]
    
    # Normal text style
    normal = styles['Normal']
    normal_font = normal.font
    normal_font.name = 'Arial'
    normal_font.size = Pt(11)


def parse_markdown_table(table_text: str) -> list[list[str]]:
    """
    Parse a Markdown table into a 2D list.
    Format: | Header1 | Header2 |
            |---------|---------|
            | Data1   | Data2   |
    """
    lines = [line.strip() for line in table_text.strip().split('\n') if line.strip()]
    table_data = []
    
    for line in lines:
        # Skip separator lines (|---|---| etc.) - include | in char class for multi-column tables
        if re.match(r'^\|[\s\-:|]+\|$', line.replace(' ', '')):
            continue
        
        # Parse data rows
        if line.startswith('|') and line.endswith('|'):
            cells = [cell.strip() for cell in line[1:-1].split('|')]
            if cells:
                table_data.append(cells)
    
    return table_data


def add_table_to_doc(doc: Document, table_data: list[list[str]], report_level: str):
    """Add a formatted table to the document with level-appropriate styling."""
    if not table_data or len(table_data) < 2:
        return
    
    # Define styling by level - use tuples (r, g, b) for shading colors
    level_styles = {
        "student": {
            "header_bg": (52, 152, 219),   # Bright blue
            "header_text": RGBColor(255, 255, 255),
            "alt_row": (235, 245, 251)     # Light blue
        },
        "professor": {
            "header_bg": (52, 73, 94),     # Professional slate
            "header_text": RGBColor(255, 255, 255),
            "alt_row": (236, 240, 241)     # Light gray
        },
        "researcher": {
            "header_bg": (33, 33, 33),     # Academic dark
            "header_text": RGBColor(255, 255, 255),
            "alt_row": (245, 245, 245)     # Very light gray
        }
    }
    
    style = level_styles.get(report_level, level_styles["student"])
    
    # Create table
    table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
    table.style = 'Light Grid Accent 1'
    
    # Import XML helpers once
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml
    
    # Format cells
    for row_idx, row in enumerate(table_data):
        for col_idx, cell_text in enumerate(row):
            cell = table.rows[row_idx].cells[col_idx]
            cell.text = cell_text
            
            # Format header row
            if row_idx == 0:
                cell_paragraph = cell.paragraphs[0]
                cell_paragraph.runs[0].font.bold = True
                cell_paragraph.runs[0].font.color.rgb = style["header_text"]
                cell_paragraph.runs[0].font.size = Pt(10)
                
                # Set background color using tuple
                r, g, b = style["header_bg"]
                shading_elm = parse_xml(r'<w:shd {} w:fill="{:02X}{:02X}{:02X}"/>'.format(
                    nsdecls('w'), r, g, b
                ))
                cell._element.get_or_add_tcPr().append(shading_elm)
            else:
                # Alternate row colors
                if row_idx % 2 == 0:
                    r, g, b = style["alt_row"]
                    shading_elm = parse_xml(r'<w:shd {} w:fill="{:02X}{:02X}{:02X}"/>'.format(
                        nsdecls('w'), r, g, b
                    ))
                    cell._element.get_or_add_tcPr().append(shading_elm)
                
                cell.paragraphs[0].runs[0].font.size = Pt(10)
    
    doc.add_paragraph()  # Spacing after table


def process_markdown_content(doc: Document, content: str, report_level: str):
    """
    Process content that may contain Markdown formatting including tables.
    Handles: **bold**, *italic*, [links](url), tables, and basic formatting.
    """
    # More robust table detection: find tables by looking for header + separator pattern
    # then greedily consume all following lines that look like table rows
    lines = content.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i].strip()
        
        # Check if this could be the start of a table (line with | characters)
        if line.startswith('|') and line.endswith('|') and line.count('|') >= 2:
            # Look ahead for separator line
            if i + 1 < len(lines):
                next_line = lines[i + 1].strip()
                # Check if it's a separator (|---|---|)
                if re.match(r'^\|[\s\-:|]+\|$', next_line.replace(' ', '')):
                    # Found a table! Collect all table rows
                    table_lines = [line, next_line]
                    j = i + 2
                    while j < len(lines):
                        row_line = lines[j].strip()
                        # Continue if line starts with | (table row)
                        if row_line.startswith('|'):
                            # Ensure it ends with | or add it
                            if not row_line.endswith('|'):
                                row_line += ' |'
                            table_lines.append(row_line)
                            j += 1
                        elif row_line == '':
                            # Empty line might be part of table, check next
                            j += 1
                        else:
                            # Non-table line, stop
                            break
                    
                    # Parse and render the table
                    table_text = '\n'.join(table_lines)
                    table_data = parse_markdown_table(table_text)
                    if table_data and len(table_data) > 1:
                        add_table_to_doc(doc, table_data, report_level)
                    
                    i = j
                    continue
        
        # Not a table - process as regular content
        if line:
            # Check for image markdown: ![caption](url)
            img_match = re.match(r'^!\[([^\]]*)\]\(([^\)]+)\)$', line)
            if img_match:
                caption = img_match.group(1)
                img_url = img_match.group(2)
                try:
                    # Download image from URL
                    response = requests.get(img_url, timeout=15, headers={
                        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
                    })
                    if response.status_code == 200:
                        img_data = io.BytesIO(response.content)
                        # Add image to document (width 5 inches)
                        try:
                            doc.add_picture(img_data, width=Inches(5))
                            # Add caption below image
                            if caption:
                                caption_para = doc.add_paragraph()
                                caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                caption_run = caption_para.add_run(f"Fig: {caption}")
                                caption_run.italic = True
                                caption_run.font.size = Pt(10)
                            doc.add_paragraph()  # Space after image
                            print(f"  📷 Added image to DOCX: {caption or img_url[:50]}")
                        except Exception as img_err:
                            print(f"  ⚠️ Failed to add image to DOCX: {img_err}")
                    else:
                        print(f"  ⚠️ Failed to download image (HTTP {response.status_code}): {img_url[:50]}")
                except Exception as e:
                    print(f"  ⚠️ Failed to include image in DOCX: {e}")
                i += 1
                continue
            
            # Handle bullet points
            if line.startswith('- ') or line.startswith('• '):
                para_text = line[2:]
                para = doc.add_paragraph(style='List Bullet')
                add_formatted_text(para, para_text)
            elif re.match(r'^\d+\.\s', line):
                # Numbered list item
                num_match = re.match(r'^\d+\.\s+(.*)', line)
                if num_match:
                    para_text = num_match.group(1)
                else:
                    para_text = line
                para = doc.add_paragraph(style='List Number')
                add_formatted_text(para, para_text)
            elif line.startswith('#'):
                # Heading
                heading_match = re.match(r'^(#+)\s*(.*)', line)
                if heading_match:
                    level = min(len(heading_match.group(1)), 3)  # Max heading level 3
                    text = heading_match.group(2)
                    doc.add_heading(text, level=level)
            else:
                # Regular paragraph
                para = doc.add_paragraph()
                add_formatted_text(para, line)
        
        i += 1


def add_formatted_text(paragraph, text: str):
    """Add text with Markdown formatting (**bold**, *italic*) to a paragraph.
    
    Processes both bold and italic patterns together to handle cases where
    they appear in any order throughout the text.
    """
    import re
    
    # Combined pattern: matches **bold** or *italic* (non-greedy)
    # Process bold first to avoid confusion with italic markers
    combined_pattern = r'(\*\*(.+?)\*\*|\*([^*]+?)\*)'
    
    current_pos = 0
    
    for match in re.finditer(combined_pattern, text):
        # Add normal text before this match
        if match.start() > current_pos:
            paragraph.add_run(text[current_pos:match.start()])
        
        full_match = match.group(0)
        
        if full_match.startswith('**') and full_match.endswith('**'):
            # Bold text
            bold_text = match.group(2)
            if bold_text:
                run = paragraph.add_run(bold_text)
                run.bold = True
        elif full_match.startswith('*') and full_match.endswith('*'):
            # Italic text
            italic_text = match.group(3)
            if italic_text:
                run = paragraph.add_run(italic_text)
                run.italic = True
        
        current_pos = match.end()
    
    # Add any remaining text after the last match
    if current_pos < len(text):
        paragraph.add_run(text[current_pos:])


def add_hyperlink(paragraph, url: str, text: str):
    """Add a hyperlink to a paragraph."""
    # This is a simplified version - full hyperlink implementation requires more XML manipulation
    run = paragraph.add_run(text)
    run.font.color.rgb = RGBColor(0, 0, 255)
    run.font.underline = True
    # Note: For full hyperlink functionality, you'd need to add relationship XML


@tool(args_schema=ExportInput)
def export_to_docx(
    sections: list[DocSection],
    filename: str,
    report_level: str = "student",
    title: str = "Research Report",
    config: RunnableConfig = None,
) -> str:
    """
    Exports structured report sections to a professional .docx file with level-appropriate formatting.
    
    Report Levels:
    - student: Friendly, readable, clear (blue theme)
    - professor: Professional, organized (slate theme)
    - researcher: Formal, academic (dark theme)
    """
    
    # Strip leading slashes to prevent issues
    clean_filename = filename.lstrip("/\\")
    if not clean_filename.lower().endswith('.docx'):
        clean_filename += '.docx'
    
    # Create document
    doc = Document()
    
    # Setup level-specific styles
    setup_styles(doc, report_level)
    
    # Add cover page
    title_paragraph = doc.add_paragraph()
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_paragraph.add_run(title)
    title_run.font.size = Pt(24 if report_level == "student" else 22)
    title_run.font.bold = True
    
    # Add level indicator
    doc.add_paragraph()
    level_names = {
        "student": "Educational Guide",
        "professor": "Teaching Resource",
        "researcher": "Research Report"
    }
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run(level_names.get(report_level, "Report"))
    subtitle_run.font.size = Pt(14)
    subtitle_run.font.italic = True
    
    # Add date
    from datetime import datetime
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_para.add_run(f"Generated: {datetime.now().strftime('%B %d, %Y')}")
    date_run.font.size = Pt(11)
    
    # Page break after cover
    doc.add_page_break()
    
    # Add sections
    for section in sections:
        # Add section heading
        doc.add_heading(section.heading, level=1)
        
        # Process content with Markdown support
        # This supports inline images via ![caption](url) markdown syntax
        process_markdown_content(doc, section.content, report_level)
        
        # Add spacing between sections
        doc.add_paragraph()
    
    # Save to in-memory buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    doc_data = buffer.read()
    
    # Encode as base64 for transmission to frontend
    doc_base64 = base64.b64encode(doc_data).decode('utf-8')
    doc_size_kb = len(doc_data) / 1024
    
    # Store download data for backend recovery
    download_data = {"filename": clean_filename, "data": doc_base64}
    _store_pending_download(download_data)
    
    print(f"  📄 DOCX generated: {clean_filename} ({doc_size_kb:.1f} KB)")
    
    return f'[DOWNLOAD_DOCX]{{"filename": "{clean_filename}", "data": "{doc_base64}"}}'